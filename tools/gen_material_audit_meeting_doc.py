# -*- coding: utf-8 -*-
"""生成《材料审核主持料控投料专题会》Word 文档。

用途：材料审核岗以主持人身份召开料控投料异常专题会的完整主持稿。
原则：不编造任何业务数字，所有需现场填报的数值一律以【 】占位。
"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn


# ---------- 基础样式工具 ----------

def set_run(run, name="宋体", size=10.5, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)


def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    set_run(p.add_run(text), name="黑体", size=20, bold=True)
    return p


def add_subtitle(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(16)
    set_run(p.add_run(text), name="楷体", size=12, color=(0x59, 0x59, 0x59))
    return p


def add_h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(8)
    set_run(p.add_run(text), name="黑体", size=14, bold=True, color=(0x1F, 0x3B, 0x73))
    return p


def add_h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    set_run(p.add_run(text), name="黑体", size=12, bold=True)
    return p


def add_body(doc, text, indent=True, size=10.5):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(4)
    if indent:
        p.paragraph_format.first_line_indent = Pt(size * 2)
    set_run(p.add_run(text), size=size)
    return p


def add_bullet(doc, text, size=10.5, level=0):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.4
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(0.75 + 0.6 * level)
    p.paragraph_format.first_line_indent = Cm(-0.5)
    set_run(p.add_run("— " if level else "• "), size=size)
    set_run(p.add_run(text), size=size)
    return p


def add_script(doc, text):
    """主持逐字稿（可直接照读）。"""
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.right_indent = Cm(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    set_run(p.add_run("【主持口径】"), name="黑体", size=10, bold=True,
            color=(0xC0, 0x50, 0x4D))
    set_run(p.add_run(text), name="楷体", size=11, color=(0x33, 0x33, 0x33))
    return p


def add_table(doc, header, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(header))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, h in enumerate(header):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run(p.add_run(h), name="黑体", size=9.5, bold=True)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.paragraph_format.line_spacing = 1.25
            set_run(p.add_run(str(val)), size=9)
    if widths:
        for r in t.rows:
            for i, w in enumerate(widths):
                r.cells[i].width = Cm(w)
    return t


# ---------- 正文构建 ----------

def build(path):
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(2.2)
    sec.bottom_margin = Cm(2.2)
    sec.left_margin = Cm(2.6)
    sec.right_margin = Cm(2.4)

    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style.font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    add_title(doc, "投料异常专题会 · 主持方案")
    add_subtitle(doc, "主持方：材料审核    参会方：料控（主责）／生产／仓储／工艺")

    # ===== 零、使用说明 =====
    add_h1(doc, "零、本文档使用说明（会前 10 分钟读完）")
    add_body(doc, "本文档是材料审核岗主持本次会议的完整脚本，含开场定调、口径说明、"
                  "分类通报、逐类提问清单、现场决议表与会后跟踪机制，可直接照稿主持。")
    add_bullet(doc, "凡出现【 】的位置，均为需你在会前从 ZPP011 主表／看板导出后自行填写的"
                    "实际数值，本文档不预设任何业务数字。")
    add_bullet(doc, "标注「主持口径」的段落为可直接朗读的话术，措辞已按"
                    "「审核提问、不做定性」的分寸校准。")
    add_bullet(doc, "会议定位是审核通报＋责任确认，不替料控做管理决策；"
                    "所有整改动作均以料控／生产签认责任人和时限的形式落地。")

    # ===== 一、会议基本信息 =====
    add_h1(doc, "一、会议基本信息")
    add_table(
        doc,
        ["项目", "内容"],
        [
            ["会议名称", "投料异常专题会（材料审核通报＋整改责任确认）"],
            ["时间／地点", "2026 年【  】月【  】日  【  】:【  】    地点：【      】"],
            ["主持人", "材料审核  【      】"],
            ["必到人员", "料控主管、投料现场负责人、仓储配料、工艺／BOM 维护、生产计划"],
            ["数据区间", "【      】年【   】月  至  【      】年【   】月"],
            ["数据来源", "ZPP011 偏差分析导出主表、负损看板、偏差率预警看板、隔离区台账"],
            ["会议时长", "建议 60 分钟（通报 15 / 质询 25 / 决议 15 / 收口 5）"],
        ],
        widths=[3.2, 12.0],
    )

    # ===== 二、会议目标 =====
    add_h1(doc, "二、会议目标（开场必须讲明，避免被理解为追责会）")
    add_bullet(doc, "第一，摊开审核口径。把「什么样的记录会被判为异常」当众讲清，"
                    "让后续每一条讨论都建立在同一套判定规则上。")
    add_bullet(doc, "第二，分类通报本期异常。按性质分成四类，每类分别对应不同的管理原因，"
                    "不混为一谈。")
    add_bullet(doc, "第三，逐类确认责任归属与整改时限。产出可跟踪的决议清单，"
                    "而不是「大家注意一下」。")
    add_bullet(doc, "第四，明确下期复核方式。本期未闭环项进入下月审核跟踪台账。")

    add_script(doc, "今天这个会不是来追谁的责任，是来把账对齐。我先把审核判定的口径"
                    "完整念一遍，大家听完如果觉得规则本身有问题，当场提，我们改规则；"
                    "如果规则没问题，那按规则筛出来的记录就得逐条给出解释。"
                    "最后我们把需要整改的项定人定时间，会后我按台账跟踪。")

    # ===== 三、审核口径 =====
    add_h1(doc, "三、审核判定口径说明（先立规矩，再谈数据）")
    add_body(doc, "本次异常记录全部由 ZPP011 偏差分析工具按下列规则自动筛出，"
                  "非人工主观挑选。口径公开的目的是让质询有共同基准。")

    add_h2(doc, "3.1 四类异常的判定规则")
    add_table(
        doc,
        ["类别", "判定规则", "业务含义", "本期条数"],
        [
            ["A 类  非耗用",
             "数量-实际 = 0  且  数量-定额 > 0（偏差率恒为 −100%）",
             "计划要求投料，但实际耗用记录为零，即一克未动",
             "【    】条"],
            ["B 类  负损少投",
             "0 ≤ 实际 < 定额（偏差为负但非零耗用）",
             "投了但未投足，包装类与辅料为高发区",
             "【    】条"],
            ["C 类  计量口径",
             "定额与实际计量单位不一致（如 kg 对 张／个／卷）",
             "统计上表现为偏差，实为单位换算口径问题",
             "【    】条"],
            ["D 类  疑难隔离",
             "已人工移入隔离区并填写疑难原因的记录",
             "系统判不清、需人工结论的反复出现项",
             "【    】条"],
        ],
        widths=[2.4, 5.0, 5.0, 2.0],
    )

    add_h2(doc, "3.2 已排除项（必须当众声明，否则会被质疑口径不公）")
    add_body(doc, "A 类中已完成替代料配对的记录，本次不列入质询范围。材料审核已在会前"
                  "逐条完成替代料配对识别，确认属于「使用代用料、原料号显示零耗用」的情形，"
                  "该部分视为已闭环。")
    add_script(doc, "先讲清楚一件事：非耗用里面，凡是能对上替代料配对的，我这边已经全部挑出来了，"
                    "今天不算在问题里。也就是说，今天摆在桌上的非耗用记录，"
                    "都是替代料解释不了的那一批。")

    # ===== 四、分类通报与质询 =====
    add_h1(doc, "四、分类通报与质询（会议主体，按 A→B→C→D 顺序）")

    # A 类
    add_h2(doc, "4.1 A 类｜非耗用（已剔除替代料配对）— 本次会议第一重点")
    add_body(doc, "剔除替代料后仍挂在「定额有要求、实际耗用为零」上的记录，只剩两种可能，"
                  "且性质完全不同，必须逐条区分：")
    add_bullet(doc, "情形一：BOM／配方临时变更，定额未同步。属流程问题，"
                    "责任落在工艺变更与主数据维护环节。")
    add_bullet(doc, "情形二：实际漏投。属现场执行问题，意味着该批次实际配方与标准不符，"
                    "同时构成成本失真与质量风险，是本次审核最需要压实的一项。")
    add_body(doc, "质询清单（逐条要求答复，不接受「回去查一下」作为最终答复，"
                  "只接受「回去查、谁查、什么时候给结论」）：", indent=False)
    add_table(
        doc,
        ["序", "质询问题", "应答方", "现场答复"],
        [
            ["1", "该物料本批次是否发生 BOM／配方变更？变更单号是多少？", "工艺／BOM 维护", ""],
            ["2", "若已变更，定额为何未同步？主数据同步的触发节点是谁负责？", "工艺／料控", ""],
            ["3", "若未变更，则该料实际是否投入？现场投料记录如何佐证？", "投料现场", ""],
            ["4", "若确认漏投，该批次产品是否已放行？是否需追溯质量判定？", "生产／质量", ""],
            ["5", "同一物料是否在多个批次重复出现同类问题？（判断偶发还是系统性）", "材料审核提出", ""],
        ],
        widths=[1.0, 7.6, 3.0, 3.2],
    )
    add_script(doc, "这一批记录的共同点是：系统里这个料的实际耗用是零，而定额要求它必须投。"
                    "替代料的可能我已经排掉了。所以只有两个答案——要么配方改了定额没跟着改，"
                    "要么就是真没投。我需要的不是解释情绪，是每一条对应到哪一种，"
                    "以及能不能拿出单据。")

    # B 类
    add_h2(doc, "4.2 B 类｜负损少投（包装类与辅料高发）")
    add_body(doc, "负损指投了但未投足。从看板监控看，物料名称含彩罐、托盘、手包袋等"
                  "包装及辅料类是高发区。原因需分开归类，不可一刀切追责：")
    add_bullet(doc, "缺料导致少投——属供应与配送节奏问题，责任在仓储／计划。")
    add_bullet(doc, "称量误差与损耗未补——属现场操作与补料规则问题。")
    add_bullet(doc, "计量口径不一致造成的「假负损」——需转入 C 类处理，"
                    "不应计入现场责任。")
    add_table(
        doc,
        ["序", "质询问题", "应答方", "现场答复"],
        [
            ["1", "本期负损记录中，属真实少投与属口径问题的各占多少？", "料控", ""],
            ["2", "真实少投中，因缺料造成的占比多少？缺料是否有预警记录？", "仓储／计划", ""],
            ["3", "现场少投是否有补料机制？补料未执行的原因是什么？", "投料现场", ""],
            ["4", "包装类物料的允差标准是否明确？现行允差是否合理？", "工艺／料控", ""],
        ],
        widths=[1.0, 7.6, 3.0, 3.2],
    )
    add_script(doc, "负损这块我要提醒一句：里面肯定有一部分不是现场的错，是计量口径对不上"
                    "算出来的假负损。所以我不会拿总条数去压现场。但请料控今天给我一个拆分——"
                    "真少投多少条、口径问题多少条。拆不出来，说明我们的允差和单位标准本身要先修。")

    # C 类
    add_h2(doc, "4.3 C 类｜计量口径不一致（治本项，优先级高于追单条）")
    add_body(doc, "包装类物料常出现定额按重量、实际按件数记录的情况，导致统计层面产生"
                  "系统性偏差。该类问题不解决，B 类的负损数据将长期失真，"
                  "审核与料控双方都会在错误的基数上争论。")
    add_bullet(doc, "需求一：清单化列出所有定额与实际计量单位不一致的物料。")
    add_bullet(doc, "需求二：为每个物料确定唯一的换算系数并固化到主数据。")
    add_bullet(doc, "需求三：明确换算系数的维护责任人与变更流程。")
    add_script(doc, "这一类我建议今天不逐条追，因为追一条修一条没有意义。"
                    "我要的是一份口径不一致的物料清单和统一的换算规则。"
                    "这个修好了，负损的数据才可信，以后我们才有得谈。")

    # D 类
    add_h2(doc, "4.4 D 类｜疑难隔离项（信息价值最高，最值得逐条过）")
    add_body(doc, "隔离区中人工填写的疑难原因，是系统规则判不清、且反复出现的记录，"
                  "其备注原因往往比主表数字更能说明问题根源。会前已按原因归类，"
                  "本环节逐类确认结论。")
    add_table(
        doc,
        ["序", "疑难原因归类", "条数", "责任方", "本次结论／处理方式"],
        [
            ["1", "【               】", "【   】", "【      】", ""],
            ["2", "【               】", "【   】", "【      】", ""],
            ["3", "【               】", "【   】", "【      】", ""],
            ["4", "【               】", "【   】", "【      】", ""],
        ],
        widths=[1.0, 4.6, 1.4, 2.2, 5.6],
    )
    add_body(doc, "同时需说明自动规则的局限：自动已读与自动隔离规则依靠物料名称关键词"
                  "结合偏差区间粗筛，既可能漏（关键词未覆盖的物料）也可能误伤"
                  "（偏差在阈值内但性质异常）。因此疑难项的最终结论一律以人工备注为准，"
                  "自动规则仅作第一道粗网，需持续校准。", indent=True)
    add_script(doc, "隔离区这些是我们自己人工标出来的疑难，含金量最高。"
                    "顺便说明一下，自动规则是按名称关键词加偏差区间筛的，会漏也会误伤，"
                    "所以规则筛出来的结果不是结论，人工备注才是结论。"
                    "谁发现规则筛错了，直接跟我提，我来校准规则。")

    # ===== 五、现场决议 =====
    add_h1(doc, "五、现场决议与责任分工（会议产出物，须当场填写并签认）")
    add_body(doc, "本表为会议唯一正式产出。未落到本表的讨论视为未形成决议，"
                  "会后不予跟踪。", indent=False)
    add_table(
        doc,
        ["序", "决议事项", "类别", "责任人", "完成时限", "验收方式"],
        [
            ["1", "", "A 类", "", "", ""],
            ["2", "", "A 类", "", "", ""],
            ["3", "", "B 类", "", "", ""],
            ["4", "", "C 类", "", "", ""],
            ["5", "", "D 类", "", "", ""],
            ["6", "", "", "", "", ""],
        ],
        widths=[1.0, 5.4, 1.4, 2.0, 2.2, 3.0],
    )

    add_h2(doc, "5.1 分工原则（有争议时按此裁定）")
    add_bullet(doc, "定额、BOM、换算系数类问题——工艺与主数据维护责任。")
    add_bullet(doc, "配送、缺料、库存供给类问题——仓储与计划责任。")
    add_bullet(doc, "投料执行、补料、记录填报类问题——投料现场责任。")
    add_bullet(doc, "规则阈值、关键词、审核口径类问题——材料审核责任，由审核方修订。")

    # ===== 六、会后跟踪 =====
    add_h1(doc, "六、会后跟踪机制")
    add_bullet(doc, "会议结束当日，将决议表转为审核跟踪台账，逐项标注状态"
                    "（未启动／进行中／已闭环）。")
    add_bullet(doc, "下期审核时，先复核上期未闭环项，再通报本期新增异常；"
                    "连续两期未闭环的事项升级上报。")
    add_bullet(doc, "自动已读与自动隔离规则的校准，由材料审核根据本次反馈更新，"
                    "更新后在下次会议开场说明变更内容。")
    add_bullet(doc, "C 类口径清单完成后，重新统计 B 类负损基数，"
                    "并在下期会议对比修正前后差异，验证治本效果。")

    # ===== 七、主持人节奏控制 =====
    add_h1(doc, "七、主持节奏控制卡（现场看这一页就够）")
    add_table(
        doc,
        ["时段", "环节", "主持动作要点"],
        [
            ["0—5 分", "开场定调", "声明会议不是追责会；讲明产出物是决议表"],
            ["5—15 分", "口径宣读", "念完四类判定规则；声明替代料已剔除；征询规则异议"],
            ["15—30 分", "A 类质询", "逐条区分「配方变更未同步」与「实际漏投」；追单据"],
            ["30—40 分", "B 类质询", "要求料控当场拆分真少投与假负损"],
            ["40—45 分", "C 类立项", "不追单条，只要口径清单与换算规则责任人"],
            ["45—55 分", "D 类过账", "按归类逐项定结论；收集规则误判反馈"],
            ["55—60 分", "决议收口", "逐条复述决议、责任人、时限；确认无异议后散会"],
        ],
        widths=[2.0, 2.6, 10.6],
    )

    add_h2(doc, "7.1 现场控场备用话术")
    add_bullet(doc, "对方以「工作量大、记不清」回避时：我理解现场忙，"
                    "所以我不要求现在给结论，我要的是谁负责查、什么时候给我结论。")
    add_bullet(doc, "对方质疑数据不准时：数据是按刚才念过的规则自动筛的，"
                    "如果你认为规则不合理，请指出哪一条规则有问题，我当场记录并修订。")
    add_bullet(doc, "讨论跑偏为部门争执时：这一条先记入待定，责任归属会后由我按分工原则"
                    "出书面意见，现在先过下一条。")
    add_bullet(doc, "对方要求放宽阈值时：阈值可以谈，但需要给出依据，"
                    "比如工艺允差或历史波动区间，凭感觉调整我不接受。")

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    set_run(p.add_run("—— 本文档由材料审核岗编制，仅供本次会议主持使用 ——"),
            name="楷体", size=9, color=(0x80, 0x80, 0x80))
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(path)
    return path


if __name__ == "__main__":
    out_dir = r"E:\Users\Administrator\Desktop"
    if not os.path.isdir(out_dir):
        out_dir = os.getcwd()
    out = os.path.join(out_dir, "投料异常专题会_材料审核主持方案.docx")
    build(out)
    print("已生成:", out)
    print("大小(KB):", round(os.path.getsize(out) / 1024, 1))
